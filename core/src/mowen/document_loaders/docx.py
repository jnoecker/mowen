"""DOCX document loader using python-docx."""

from __future__ import annotations

from pathlib import Path

try:
    import docx
except ImportError:
    raise ImportError("Install python-docx: pip install mowen[docx]")

from mowen.document_loaders.base import DocumentLoader
from mowen.exceptions import DocumentLoadError
from mowen.types import Document


class DOCXLoader(DocumentLoader):
    """Loads a DOCX file as a Document."""

    @staticmethod
    def supported_extensions() -> list[str]:
        return [".docx"]

    def load(self, path: Path, author: str | None = None, title: str | None = None) -> Document:
        path = Path(path)
        if not path.exists():
            raise DocumentLoadError(f"File not found: {path}")
        try:
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except DocumentLoadError:
            raise
        except Exception as e:
            raise DocumentLoadError(f"Cannot read DOCX {path}: {e}") from e
        return Document(
            text=text,
            author=author,
            title=title or path.stem,
        )
